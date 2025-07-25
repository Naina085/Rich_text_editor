
#upload MS word document


# from flask import Flask, request, jsonify
# from flask_cors import CORS  # 👈 Add this
# from docx import Document
# import base64, os
# from werkzeug.utils import secure_filename

# app = Flask(__name__)
# CORS(app)  # 👈 This enables CORS for all routes

# UPLOAD_FOLDER = 'uploads'
# os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# def extract_docx_content(docx_path):
#     doc = Document(docx_path)
#     html = ""
#     for para in doc.paragraphs:
#         html += f"<p>{para.text}</p>"
#     for rel in doc.part._rels:
#         rel = doc.part._rels[rel]
#         if "image" in rel.target_ref:
#             img_data = rel.target_part.blob
#             img_base64 = base64.b64encode(img_data).decode("utf-8")
#             img_tag = f'<img src="data:image/png;base64,{img_base64}" style="max-width:100%;"/><br>'
#             html += img_tag
#     return html

# @app.route('/upload', methods=['POST'])
# def upload():
#     if 'docx' not in request.files:
#         return jsonify({'error': 'No file'}), 400

#     file = request.files['docx']
#     filename = secure_filename(file.filename)
#     file_path = os.path.join(UPLOAD_FOLDER, filename)
#     file.save(file_path)

#     html_content = extract_docx_content(file_path)
#     return jsonify({'html_content': html_content})

# if __name__ == '__main__':
#     app.run(debug=True)


from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from werkzeug.utils import secure_filename
import base64, os

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def extract_images_map(doc):
    """Returns a dict mapping rId -> base64 image tag"""
    image_map = {}
    for rel in doc.part.rels.values():
        if rel.reltype == RT.IMAGE:
            rid = rel.rId
            img_data = rel.target_part.blob
            img_ext = rel.target_ref.split('.')[-1]
            img_base64 = base64.b64encode(img_data).decode('utf-8')
            img_tag = f'<img src="data:image/{img_ext};base64,{img_base64}" style="max-width:100%;"/>'
            image_map[rid] = img_tag
    return image_map

def run_to_html(run, image_map):
    if "graphic" in run._element.xml and "r:embed" in run._element.xml:
        # Try to extract image relationship ID
        for rid in image_map:
            if f'r:embed="{rid}"' in run._element.xml:
                return image_map[rid]

    text = run.text.replace("\n", "<br>")
    if not text:
        return ""

    style = ""
    if run.font.color and run.font.color.rgb:
        style += f"color:#{run.font.color.rgb};"
    if run.font.size:
        pt = run.font.size.pt
        style += f"font-size:{pt}pt;"

    html = text
    if run.bold:
        html = f"<b>{html}</b>"
    if run.italic:
        html = f"<i>{html}</i>"
    if run.underline:
        html = f"<u>{html}</u>"

    return f'<span style="{style}">{html}</span>' if style else html

def is_list(paragraph):
    return paragraph.style.name.lower().startswith(("list", "bullet", "number"))

def paragraph_to_html(paragraph, image_map):
    html = "".join(run_to_html(run, image_map) for run in paragraph.runs)
    tag = "p"
    style = paragraph.style.name.lower()

    if "heading" in style:
        try:
            level = int(style.replace("heading", "").strip())
            if 1 <= level <= 6:
                tag = f"h{level}"
        except:
            pass

    return f"<{tag}>{html}</{tag}>"

def table_to_html(table, image_map):
    html = "<table border='1' style='border-collapse:collapse;width:100%;'>"
    for row in table.rows:
        html += "<tr>"
        for cell in row.cells:
            cell_html = "".join(paragraph_to_html(p, image_map) for p in cell.paragraphs)
            html += f"<td>{cell_html}</td>"
        html += "</tr>"
    html += "</table><br>"
    return html

def extract_docx_content(docx_path):
    doc = Document(docx_path)
    image_map = extract_images_map(doc)

    html = ""
    list_buffer = []
    current_list_type = None

    def flush_list():
        nonlocal html, list_buffer, current_list_type
        if list_buffer:
            tag = "ul" if current_list_type == "bullet" else "ol"
            html += f"<{tag}>" + "".join(f"<li>{item}</li>" for item in list_buffer) + f"</{tag}>"
            list_buffer = []
            current_list_type = None

    for block in doc.element.body.iterchildren():
        if block.tag == qn("w:p"):
            paragraph = next(p for p in doc.paragraphs if p._p == block)
            if is_list(paragraph):
                list_type = "bullet" if "bullet" in paragraph.style.name.lower() else "number"
                content = "".join(run_to_html(run, image_map) for run in paragraph.runs)

                if current_list_type != list_type:
                    flush_list()
                    current_list_type = list_type

                list_buffer.append(content)
            else:
                flush_list()
                html += paragraph_to_html(paragraph, image_map)

        elif block.tag == qn("w:tbl"):
            flush_list()
            table = next(t for t in doc.tables if t._tbl == block)
            html += table_to_html(table, image_map)

    flush_list()

    return f'<div contenteditable="true" style="padding:10px;">{html}</div>'

@app.route('/upload', methods=['POST'])
def upload():
    if 'docx' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['docx']
    filename = secure_filename(file.filename)
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(file_path)

    html_content = extract_docx_content(file_path)
    return jsonify({'html_content': html_content})

if __name__ == '__main__':
    app.run(debug=True)
